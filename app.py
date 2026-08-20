import io
import re
import time
import requests
from bs4 import BeautifulSoup
import streamlit as st
from google import genai
from google.genai import types
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ----------------------------------------------------
# 1. 페이지 기본 설정
# ----------------------------------------------------
st.set_page_config(
    page_title="삼성생명 기획팀 동향분석 Agent",
    page_icon="📊",
    layout="wide"
)

# ----------------------------------------------------
# 2. 기사 본문 자동 추출 크롤러 유틸리티
# ----------------------------------------------------
def extract_article_content(user_text):
    url_pattern = re.compile(r'https?://[^\s]+')
    urls = url_pattern.findall(user_text.strip())
    
    if not urls:
        return user_text, None

    target_url = urls[0]
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
    }
    
    try:
        res = requests.get(target_url, headers=headers, timeout=8)
        res.raise_for_status()
        res.encoding = res.apparent_encoding
        
        soup = BeautifulSoup(res.text, 'html.parser')
        for tag in soup(['script', 'style', 'header', 'footer', 'nav', 'aside', 'iframe']):
            tag.decompose()
            
        naver_news = soup.find('article', id='dic_area') or soup.find('div', id='newsct_article') or soup.find('div', id='articeBody')
        if naver_news:
            return naver_news.get_text(separator='\n', strip=True), target_url
            
        article_tag = soup.find('article') or soup.find('div', class_=re.compile(r'article|news|content|body', re.I))
        if article_tag:
            paragraphs = [p.get_text(strip=True) for p in article_tag.find_all(['p', 'div']) if len(p.get_text(strip=True)) > 20]
            if paragraphs:
                return '\n'.join(paragraphs), target_url

        text_lines = [p.get_text(strip=True) for p in soup.find_all('p') if len(p.get_text(strip=True)) > 20]
        if text_lines:
            return '\n'.join(text_lines), target_url
            
        return user_text, target_url
    except Exception:
        return user_text, target_url

# ----------------------------------------------------
# 3. Gem 시스템 지침 (실시간 심층 추가 검색 활성화)
# ----------------------------------------------------
SYSTEM_INSTRUCTION = """
# Role
당신은 삼성생명 기획팀의 "동향분석 및 보고서 작성 전문 Agent"입니다.
국내 1위 생명보험사인 삼성생명의 기획팀 직원 관점에서 전달된 메인 기사를 분석하고, 실시간 추가 구글 검색(Deep Dive Search)을 통해 배경 맥락과 최신 비교 데이터를 덧붙여 정교한 사내 보고서를 작성합니다.

# Core Workflow (4단계 보고서 작성 프로세스)
[Step 1. 메인 기사 분석]
- 제공된 기사 원문(또는 텍스트)을 분석하여 주요 사실관계, 핵심 수치, 이슈 맥락을 정확히 파악합니다.

[Step 2. 심층 검색 및 데이터 결합 (Deep Dive Search)]
- 제공된 메인 기사 외에 반드시 **구글 검색 도구(Google Search)를 능동적으로 실행**하여 관련 최신 데이터를 실시간으로 보완·학습합니다:
  1) 메인 이슈 관련 최신 업계 동향 및 금융당국 규제/제도 수치 (CSM, K-ICS 비율, 할인율, 무·저해지 해약률 등)
  2) 타 생보사(한화, 교보 등) 대응 현황 및 삼성생명 관련 사업/재무 비교 데이터 (CSM 잔액/신계약, 순이익, K-ICS, FC 규모 등)

[Step 3. 삼성생명 맞춤형 전략적 시사점 도출]
- 입력된 기사의 핵심 내용 및 추가 검색된 시장 환경에 입각하여, 삼성생명의 전략적 영향과 대응 방향('So What')을 자율적·논리적으로 해석하여 도출합니다.

[Step 4. 규격화된 보고서 작성]
- 아래 Output Structure & Formatting Rules를 준수하여 보고서를 작성합니다.

# Output Structure & Formatting Rules (출력 및 작성 규격)
1. 문체 및 특수기호 위계 구조 원칙 (엄격 준수):
   - [Page 1], [Page 2] 같은 라벨이나 <표 1> 표는 일체 작성하지 않습니다.
   - 대항목 (네모): `□` -> 대주제 및 핵심 테마 명시 (반드시 `□ [요약]` 및 `□ [시사점]` 표기 준수)
   - 중항목 (찍): `-` -> 반드시 하위 소항목(·)들의 내용을 포괄하는 간결하고 압축적인 '한 문장 요약/헤드라인' 형태로 작성하며 명사형 종결
   - 소항목 (땡): `·` -> 상위 중항목(-)을 뒷받침하는 구체적인 세부 팩트, 통계 수치, 실행 과제, 메커니즘을 상세히 서술
   - 문체: 명확하고 격식 있는 보고서용 개조식 명사형 종결문 (~확대, ~추진, ~견지, ~구축, ~달성, ~확보 등)
   - [문장 길이]: 바탕체 15pt 기준 각 항목(□, -, ·)의 문장이 한 줄에 다 들어오거나 최대 2줄을 넘지 않도록 서술

2. 보고서 본문 구성:
---
□ [요약] (분석된 실제 기사의 핵심 제목)
  - (기사의 팩트와 추가 검색된 수치를 포괄하는 핵심 사실관계 한 문장 요약 명사형 종결)
    · 메인 기사 주요 사실관계 및 일자/배경 세부 내용
    · 삼성생명 및 업계 관련 세부 통계/지표 데이터 (추가 검색 데이터 결합)
  - (시장 변화 및 타사 동향을 포괄하는 업계 판도 한 문장 요약 명사형 종결)
    · 경쟁사 대응 현황 및 시장 판매/영업 전략 동향
    · 금융당국 규제/제도 기조 및 시장 환경 영향 분석

□ [시사점] (해당 이슈 기반 삼성생명 맞춤형 전략 방향 테마)
  - (기사 이슈가 삼성생명에 미치는 직접적 영향 및 핵심 대응 과제 한 문장 요약 명사형 종결)
    · 해당 이슈 관련 삼성생명의 사업/영업/상품/자본 측면 세부 전략 과제
    · 시장 리더십 유지를 위한 차별화 실행 방안
  - (향후 리스크 관리 및 기회 선점을 위한 중장기 전략 한 문장 요약 명사형 종결)
    · 환경 변화에 따른 선제적 리스크 통제 및 운영 최적화 방안
    · 미래 경쟁력 강화를 위한 연계 시너지 창출 방안

# 참고 자료 및 심층 출처 리스트 (References & Deep Dive Data)

1. [메인 기사 출처]
  - 출처명: 기사 제목
    · 핵심 요약: 1~2줄 컴팩트 요약

2. [심층 배경 및 비교 데이터]
  - 주요 지표 및 제도적 맥락 (실시간 검색으로 파악한 데이터)
    · 핵심 데이터: 1~2줄 핵심 수치 요약

3. [기획팀 종합 평가 및 향후 모니터링 포인트]
  - 시장 영향도 평가: 기사 이슈가 업계 및 1위사에 미치는 파급력
  - 리스크 선제 대응: 중점 모니터링 필요 요소 및 리스크 관리 방향
  - 전략적 실행 제언: 기획팀 차원의 단기/중장기 핵심 추진 과제
"""

# ----------------------------------------------------
# 4. Word 문서 생성 유틸리티 (바탕체, 15pt 서식 적용)
# ----------------------------------------------------
def set_font_style(run, name="바탕체", size_pt=15, bold=False, color_rgb=None):
    run.font.name = name
    run._r.get_or_add_rPr().set(qn('w:rFonts'), qn('w:eastAsia'))
    run._r.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color_rgb:
        run.font.color.rgb = color_rgb

def create_docx(text_content):
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    lines = text_content.split("\n")
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
            
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(line.replace("# ", "").strip())
            set_font_style(run, name="바탕체", size_pt=16, bold=True, color_rgb=RGBColor(0, 51, 102))
        elif line.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(line.replace("## ", "").strip())
            set_font_style(run, name="바탕체", size_pt=15, bold=True)
        elif stripped.startswith("□") or stripped.startswith("ㅁ"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.25
            run = p.add_run(stripped)
            set_font_style(run, name="바탕체", size_pt=15, bold=True, color_rgb=RGBColor(0, 51, 102))
        elif stripped.startswith("-"):
            p = doc.add_paragraph(stripped)
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.25
            run = p.add_run(stripped)
            set_font_style(run, name="바탕체", size_pt=15, bold=False)
        elif stripped.startswith("·") or stripped.startswith("."):
            p = doc.add_paragraph(stripped)
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.25
            run = p.add_run(stripped)
            set_font_style(run, name="바탕체", size_pt=15, bold=False)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.25
            run = p.add_run(stripped)
            set_font_style(run, name="바탕체", size_pt=15, bold=False)
                
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

# ----------------------------------------------------
# 5. 부서원 비밀번호 인증
# ----------------------------------------------------
if "authenticated" not in st.session_state:
    st.session_state["authenticated"] = False

if not st.session_state["authenticated"]:
    st.title("🔒 기획팀 전용 분석 Agent")
    pwd = st.text_input("부서 접속 비밀번호를 입력하세요", type="password")
    correct_pwd = st.secrets.get("AUTH_PASSWORD", "1234")
    
    if st.button("접속", use_container_width=True):
        if pwd == correct_pwd:
            st.session_state["authenticated"] = True
            st.rerun()
        else:
            st.error("비밀번호가 일치하지 않습니다.")
    st.stop()

# ----------------------------------------------------
# 6. 메인 앱 화면 및 분석 실행 (크롤링 + Google Search 도구 결합)
# ----------------------------------------------------
api_key = st.secrets.get("GEMINI_API_KEY")
if not api_key:
    st.error("GEMINI_API_KEY가 설정되지 않았습니다. Secrets를 확인해주세요.")
    st.stop()

client = genai.Client(api_key=api_key)

st.title("📊 삼성생명 기획팀 동향분석 Agent")
st.caption("기사 URL 또는 본문 텍스트를 입력하면 실시간 심층 검색을 거쳐 살을 덧붙인 정규 보고서를 작성합니다.")

user_input = st.text_area(
    "분석할 기사 내용 또는 기사 URL을 입력하세요:",
    height=150,
    placeholder="예: 기사 전문을 붙여넣거나 분석할 기사의 웹 링크(URL)를 입력하세요."
)

if st.button("보고서 생성 시작", type="primary", use_container_width=True):
    if not user_input.strip():
        st.warning("분석할 기사 내용이나 URL을 입력해주세요.")
    else:
        with st.spinner("기사 본문 확인 및 실시간 심층 추가 검색(Deep Dive)을 진행 중입니다..."):
            extracted_text, detected_url = extract_article_content(user_input)
            
            if detected_url:
                prompt_content = f"다음은 입력된 URL({detected_url})의 기사 본문입니다. 이 내용을 메인 뼈대로 삼고, 구글 검색을 통해 관련 최신 수치/타사 동향을 심층 보완하여 보고서를 작성하십시오:\n\n{extracted_text}"
            else:
                prompt_content = f"다음 기사 내용을 메인 뼈대로 삼고, 구글 검색을 통해 관련 최신 수치/타사 동향을 심층 보완하여 보고서를 작성하십시오:\n\n{extracted_text}"

            response_success = False
            last_error_msg = ""

            for attempt in range(3):
                try:
                    # Google Search 도구를 다시 결합하여 실시간 추가 검색 수행
                    response = client.models.generate_content(
                        model="gemini-3.6-flash",
                        contents=prompt_content,
                        config=types.GenerateContentConfig(
                            system_instruction=SYSTEM_INSTRUCTION,
                            temperature=0.2,
                            tools=[types.Tool(google_search=types.GoogleSearch())]
                        )
                    )
                    st.session_state["last_report"] = response.text
                    response_success = True
                    break
                except Exception as e:
                    last_error_msg = str(e)
                    if any(err in last_error_msg for err in ["503", "UNAVAILABLE", "429"]):
                        time.sleep(3)
                        continue
                    else:
                        break

            if not response_success:
                st.error(f"보고서 생성 중 오류가 발생했습니다: {last_error_msg}")

# ----------------------------------------------------
# 7. 생성된 보고서 표시 및 다운로드
# ----------------------------------------------------
if "last_report" in st.session_state:
    st.divider()
    st.markdown("### 📄 작성된 동향분석 보고서")
    st.markdown(st.session_state["last_report"])
    
    st.divider()
    col1, col2 = st.columns(2)
    
    with col1:
        docx_file = create_docx(st.session_state["last_report"])
        st.download_button(
            label="📥 Word 보고서 (.docx) 다운로드",
            data=docx_file,
            file_name="삼성생명_동향분석보고서.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
        
    with col2:
        st.download_button(
            label="📥 텍스트 보고서 (.txt) 다운로드",
            data=st.session_state["last_report"].encode("utf-8"),
            file_name="삼성생명_동향분석보고서.txt",
            mime="text/plain",
            use_container_width=True
        )
